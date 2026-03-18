"""DOCX handler using python-docx."""

from docx import Document


def analyze_summary(file_path):
    doc = Document(file_path)
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        entry = {
            "index": i,
            "style": p.style.name if p.style else "Normal",
            "text": p.text[:200] + ("..." if len(p.text) > 200 else ""),
        }
        paragraphs.append(entry)

    tables = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        tables.append({"index": i, "rows": rows, "cols": cols})

    return {
        "doc_type": "docx",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def analyze_search(file_path, query):
    doc = Document(file_path)
    query_lower = query.lower()
    results = []

    for i, p in enumerate(doc.paragraphs):
        if query_lower in p.text.lower():
            results.append({
                "type": "paragraph",
                "index": i,
                "style": p.style.name if p.style else "Normal",
                "text": p.text[:300],
            })

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if query_lower in cell.text.lower():
                    results.append({
                        "type": "table_cell",
                        "table_index": ti,
                        "row": ri,
                        "col": ci,
                        "text": cell.text[:300],
                    })

    return {"query": query, "matches": results}


def apply_operations(file_path, operations):
    doc = Document(file_path)
    applied = []

    for idx, op in enumerate(operations):
        op_type = op["op"]
        try:
            if op_type == "replace_text":
                _op_replace_text(doc, op)
            elif op_type == "insert_paragraph_after":
                _op_insert_paragraph_after(doc, op)
            elif op_type == "delete_paragraph":
                _op_delete_paragraph(doc, op)
            else:
                raise ValueError(f"Unknown operation: {op_type}")
            applied.append({"index": idx, "op": op_type, "status": "ok"})
        except Exception as e:
            return {
                "success": False,
                "failed_op_index": idx,
                "failed_op": op_type,
                "error": str(e),
                "applied_before_failure": applied,
            }

    doc.save(file_path)
    return {"success": True, "applied": applied}


def _op_replace_text(doc, op):
    """Locator-based replace: requires target.paragraph_index."""
    target = op.get("target")
    if not target or "paragraph_index" not in target:
        raise ValueError("replace_text requires target.paragraph_index")

    paragraph_index = target["paragraph_index"]
    old_text = op["args"]["old_text"]
    new_text = op["args"]["new_text"]

    paragraphs = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        raise ValueError(f"Invalid paragraph index: {paragraph_index} (total: {len(paragraphs)})")

    p = paragraphs[paragraph_index]
    full_text = "".join(run.text for run in p.runs)
    if old_text not in full_text:
        raise ValueError(
            f"Text not found in paragraph {paragraph_index}: {old_text[:50]}"
        )

    _replace_in_runs(p.runs, old_text, new_text)


def _replace_in_runs(runs, old_text, new_text):
    """Replace text across runs while preserving formatting of the first run."""
    combined = ""
    run_map = []  # (run_index, char_start, char_end)
    for i, run in enumerate(runs):
        start = len(combined)
        combined += run.text
        run_map.append((i, start, len(combined)))

    idx = combined.find(old_text)
    if idx == -1:
        return

    end_idx = idx + len(old_text)

    for run_i, char_start, char_end in run_map:
        if char_end <= idx or char_start >= end_idx:
            continue
        run = runs[run_i]
        run_text = run.text
        overlap_start = max(idx, char_start) - char_start
        overlap_end = min(end_idx, char_end) - char_start
        new_run_text = run_text[:overlap_start] + run_text[overlap_end:]
        if char_start <= idx < char_end:
            new_run_text = run_text[:overlap_start] + new_text + run_text[overlap_end:]
        run.text = new_run_text


def _op_insert_paragraph_after(doc, op):
    target_index = op["target"]["paragraph_index"]
    text = op["args"]["text"]
    style = op["args"].get("style", "Normal")

    paragraphs = doc.paragraphs
    if target_index < 0 or target_index >= len(paragraphs):
        raise ValueError(f"Invalid paragraph index: {target_index} (total: {len(paragraphs)})")

    target_para = paragraphs[target_index]
    new_para = doc.add_paragraph(text, style=style)
    target_para._element.addnext(new_para._element)


def _op_delete_paragraph(doc, op):
    target_index = op["target"]["paragraph_index"]
    paragraphs = doc.paragraphs
    if target_index < 0 or target_index >= len(paragraphs):
        raise ValueError(f"Invalid paragraph index: {target_index} (total: {len(paragraphs)})")

    p = paragraphs[target_index]
    p._element.getparent().remove(p._element)
